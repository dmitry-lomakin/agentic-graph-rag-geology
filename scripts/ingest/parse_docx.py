"""Word document (.doc/.docx) -> Markdown parser.

DOCX files are parsed directly with python-docx. Legacy .doc files are first
converted to .docx via LibreOffice headless mode, then parsed the same way.

Handles: headings, tables, lists, bold/italic, and image placeholders.
"""

import logging
import subprocess
import tempfile
from pathlib import Path

from scripts.ingest.models import ParsedDocument
from scripts.utils.external_tools import find_libreoffice

logger = logging.getLogger(__name__)

LIBREOFFICE_TIMEOUT = 120  # seconds


def _convert_doc_to_docx(file_path: Path) -> Path:
    """Convert a .doc file to .docx using LibreOffice headless.

    Args:
        file_path: Path to .doc file.

    Returns:
        Path to the converted .docx file in a temp directory.

    Raises:
        RuntimeError: If LibreOffice conversion fails.
    """
    tmp_dir = tempfile.mkdtemp(prefix="doc2docx_")
    logger.debug("Converting .doc -> .docx: %s -> %s", file_path.name, tmp_dir)

    try:
        soffice = find_libreoffice()
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                tmp_dir,
                str(file_path),
            ],
            capture_output=True,
            text=True,
            timeout=LIBREOFFICE_TIMEOUT,
            check=True,
        )
    except subprocess.CalledProcessError as exc:
        raise RuntimeError(
            f"LibreOffice conversion failed for {file_path.name}: {exc.stderr}"
        ) from exc
    except (subprocess.TimeoutExpired, FileNotFoundError) as exc:
        raise RuntimeError(
            f"LibreOffice conversion failed for {file_path.name}: {exc}"
        ) from exc

    converted = Path(tmp_dir) / (file_path.stem + ".docx")
    if not converted.exists():
        raise RuntimeError(
            f"LibreOffice did not produce expected output: {converted}"
        )
    return converted


def _docx_to_markdown(file_path: Path) -> tuple[str, list[str]]:
    """Extract content from a .docx file and convert to markdown.

    Args:
        file_path: Path to .docx file.

    Returns:
        Tuple of (markdown_text, list_of_warnings).
    """
    from docx import Document
    from docx.table import Table
    from docx.oxml.ns import qn

    doc = Document(str(file_path))
    lines: list[str] = []
    warnings: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]  # strip namespace

        if tag == "p":
            _process_paragraph(element, doc, lines)
        elif tag == "tbl":
            table = Table(element, doc)
            _process_table(table, lines)

    # Extract title from first heading or core properties
    title = ""
    if doc.core_properties.title:
        title = doc.core_properties.title
    else:
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                title = para.text.strip()
                break

    markdown = "\n\n".join(lines)
    return markdown, warnings


def _process_paragraph(element, doc, lines: list[str]) -> None:
    """Convert a paragraph element to markdown."""
    from docx.oxml.ns import qn

    # Find the corresponding paragraph object
    para = None
    for p in doc.paragraphs:
        if p._element is element:
            para = p
            break

    if para is None:
        return

    text = para.text.strip()
    if not text:
        # Check for images
        for rel in para._element.iter(qn("wp:inline")):
            lines.append("[Image]")
            return
        for rel in para._element.iter(qn("wp:anchor")):
            lines.append("[Image]")
            return
        return

    style_name = para.style.name if para.style else ""

    # Headings
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.replace("Heading", "").strip())
            level = min(level, 6)
        except ValueError:
            level = 1
        lines.append(f"{'#' * level} {text}")
        return

    # List items
    numpr = para._element.find(qn("w:pPr"))
    if numpr is not None and numpr.find(qn("w:numPr")) is not None:
        lines.append(f"- {text}")
        return

    # Regular paragraph with inline formatting
    formatted = _format_runs(para)
    if formatted:
        lines.append(formatted)


def _format_runs(para) -> str:
    """Format paragraph runs with bold/italic markdown."""
    parts: list[str] = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            parts.append(f"***{text}***")
        elif run.bold:
            parts.append(f"**{text}**")
        elif run.italic:
            parts.append(f"*{text}*")
        else:
            parts.append(text)
    return "".join(parts)


def _process_table(table, lines: list[str]) -> None:
    """Convert a docx Table to markdown table format."""
    rows_data: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows_data.append(cells)

    if not rows_data:
        return

    # First row as header
    header = rows_data[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in rows_data[1:]:
        # Pad row to match header column count
        while len(row) < len(header):
            row.append("")
        lines.append("| " + " | ".join(row[:len(header)]) + " |")


def parse_docx_file(file_path: Path) -> ParsedDocument:
    """Parse a .docx file into structured markdown.

    Args:
        file_path: Path to a .docx file on disk.

    Returns:
        ParsedDocument with extracted content.
    """
    logger.info("Parsing DOCX: %s", file_path.name)
    markdown, warnings = _docx_to_markdown(file_path)

    from docx import Document
    doc = Document(str(file_path))
    title = doc.core_properties.title or ""
    if not title:
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                title = para.text.strip()
                break
    if not title:
        title = file_path.stem

    word_count = len(markdown.split())
    has_figures = "[Image]" in markdown

    return ParsedDocument(
        title=title,
        markdown=markdown,
        has_figures=has_figures,
        word_count=word_count,
        source_format="docx",
        parser_used="python-docx",
        parse_warnings=warnings,
    )


def parse_doc_file(file_path: Path) -> ParsedDocument:
    """Parse a legacy .doc file into structured markdown.

    Converts to .docx via LibreOffice first, then parses with python-docx.

    Args:
        file_path: Path to a .doc file on disk.

    Returns:
        ParsedDocument with extracted content.
    """
    logger.info("Parsing DOC (via LibreOffice conversion): %s", file_path.name)
    converted_path = _convert_doc_to_docx(file_path)

    try:
        doc = parse_docx_file(converted_path)
        doc.source_format = "doc"
        if not doc.title or doc.title == converted_path.stem:
            doc.title = file_path.stem
        return doc
    finally:
        # Clean up temp file
        converted_path.unlink(missing_ok=True)
        converted_path.parent.rmdir()
